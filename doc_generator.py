"""
Document Generator for Club 13 GMP Toolkit

Generates .docx files from the actual Club 13 .dotx templates by
surgically filling values into the existing template structure
WITHOUT altering run boundaries, formatting, or layout.

Templates:
  - PK Specification Test Record Template.dotx
  - RM Specification Test Record Template.dotx
  - SOP Temp.dotx
  - CLB003 Component Receiving Record.docx
  - CLB008 QC Release Sticker.docx
  - CLB009 QC Sampled.docx
  - CLB010 Component ID Tag.docx
"""
import io
import math
import os
import re
import tempfile
import zipfile
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn

WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _val(row, key, default=""):
    """Safely get a value from a sqlite3.Row or dict."""
    try:
        v = row[key]
        return v if v is not None else default
    except (KeyError, IndexError):
        return default


# ── .dotx opener ────────────────────────────────────────────────

def open_dotx(path):
    """Open a .dotx template by patching its content type for python-docx."""
    with open(path, "rb") as f:
        data = io.BytesIO(f.read())

    patched = io.BytesIO()
    with zipfile.ZipFile(data, "r") as zin, zipfile.ZipFile(patched, "w") as zout:
        for item in zin.infolist():
            content = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                content = content.replace(
                    b"wordprocessingml.template.main+xml",
                    b"wordprocessingml.document.main+xml",
                )
            zout.writestr(item, content)

    patched.seek(0)
    return Document(patched)


# ── Run-safe helpers ────────────────────────────────────────────

def _strip_fillin_fields(paragraph, field_values):
    """Remove FILLIN merge-field code runs from a paragraph and insert values.

    Word FILLIN fields consist of 3 consecutive runs:
      <w:r><w:fldChar fldCharType="begin"/></w:r>
      <w:r><w:instrText> FILLIN "Field Name" ...</w:instrText></w:r>
      <w:r><w:fldChar fldCharType="end"/></w:r>

    This function removes those runs and inserts a text run with the
    corresponding value from field_values dict after the label run.

    field_values maps the FILLIN prompt to the replacement text, e.g.:
      {"Component No.": "PKG-60W", "Rev. No.": "02", "Component Name": "Widget"}
    """
    p_elem = paragraph._element
    W = qn("w:r")
    runs_in_xml = list(p_elem.iterchildren(W))

    # Also look inside smartTags for runs
    for st in list(p_elem.iterchildren(qn("w:smartTag"))):
        runs_in_xml.extend(st.iterchildren(W))

    # Collect field code blocks: groups of (begin_run, instr_run, end_run)
    field_blocks = []
    i = 0
    all_children = list(p_elem)
    while i < len(all_children):
        child = all_children[i]
        if child.tag == qn("w:r"):
            fc = child.find(qn("w:fldChar"))
            if fc is not None and fc.get(qn("w:fldCharType")) == "begin":
                # Start of field — collect begin, instrText, end
                block_runs = [child]
                field_name = None
                j = i + 1
                while j < len(all_children):
                    sibling = all_children[j]
                    if sibling.tag == qn("w:r"):
                        block_runs.append(sibling)
                        # Check for instrText
                        instr = sibling.find(qn("w:instrText"))
                        if instr is not None and instr.text:
                            # Extract field name from: FILLIN "Field Name" \* MERGEFORMAT
                            import re as _re
                            m = _re.search(r'FILLIN\s+"([^"]+)"', instr.text)
                            if m:
                                field_name = m.group(1)
                        # Check for end
                        fc2 = sibling.find(qn("w:fldChar"))
                        if fc2 is not None and fc2.get(qn("w:fldCharType")) == "end":
                            field_blocks.append((block_runs, field_name))
                            i = j
                            break
                    j += 1
        i += 1

    # Remove field code runs and insert value runs
    for block_runs, field_name in field_blocks:
        # Find point to insert after (the last run before the field begin)
        insert_after = block_runs[0].getprevious()

        # Get the value
        value = ""
        if field_name and field_name in field_values:
            value = str(field_values[field_name])

        # Remove all field code runs
        for r in block_runs:
            r.getparent().remove(r)

        # Create a value run (inheriting formatting from the label run if available)
        new_r = p_elem.makeelement(qn("w:r"), {})
        new_t = new_r.makeelement(qn("w:t"), {})
        new_t.text = value
        new_t.set(qn("xml:space"), "preserve")
        new_r.append(new_t)

        # Insert after the label run
        if insert_after is not None:
            insert_after.addnext(new_r)
        else:
            p_elem.insert(0, new_r)


def _fill_after_label(paragraph, label, value):
    """Fill the first empty run after a label run, preserving all run boundaries.

    Template pattern:  [label_run] [empty_run] [empty_run] ...
    Result:            [label_run] [value_run] [empty_run] ...
    """
    runs = paragraph.runs
    label_found = False
    for run in runs:
        if label in run.text:
            label_found = True
            continue
        if label_found and run.text.strip() == "":
            run.text = str(value) if value else ""
            return True
    return False


def _replace_underscores(run, value):
    """Replace a contiguous underscore placeholder in a run with a value."""
    run.text = re.sub(r"_{3,}", str(value) if value else "", run.text, count=1)


def _set_row_cell_text(row_tr, col_idx, text):
    """Set text in a specific cell of a cloned row XML element."""
    cells = row_tr.findall(qn("w:tc"))
    if col_idx >= len(cells):
        return
    for p in cells[col_idx].findall(qn("w:p")):
        runs = p.findall(qn("w:r"))
        if runs:
            # Use existing first run — preserves its rPr
            for t in runs[0].findall(qn("w:t")):
                t.text = str(text) if text else ""
                t.set(qn("xml:space"), "preserve")
                return
            # Run exists but has no <w:t> — add one
            t_elem = runs[0].makeelement(qn("w:t"), {})
            t_elem.text = str(text) if text else ""
            t_elem.set(qn("xml:space"), "preserve")
            runs[0].append(t_elem)
            return
        # No runs at all — copy rPr from a sibling cell's run if available
        rPr_source = _find_sibling_rPr(row_tr, col_idx)
        r_elem = p.makeelement(qn("w:r"), {})
        if rPr_source is not None:
            r_elem.append(deepcopy(rPr_source))
        t_elem = r_elem.makeelement(qn("w:t"), {})
        t_elem.text = str(text) if text else ""
        t_elem.set(qn("xml:space"), "preserve")
        r_elem.append(t_elem)
        p.append(r_elem)
        return


def _find_sibling_rPr(row_tr, skip_col):
    """Find an rPr element from another cell in the same row for style copying."""
    for ci, tc in enumerate(row_tr.findall(qn("w:tc"))):
        if ci == skip_col:
            continue
        for p in tc.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                rPr = r.find(qn("w:rPr"))
                if rPr is not None:
                    return rPr
    return None


# ══════════════════════════════════════════════════════════════════
# PK / RM SPECIFICATION TEST RECORD
# ══════════════════════════════════════════════════════════════════

def generate_pk_spec_record(template_path, spec, parameters,
                            completion_fields=None, attachment_paths=None):
    """Generate a PK Specification Test Record with spec info filled in."""
    return _generate_spec_record(template_path, spec, parameters,
                                 completion_fields=completion_fields,
                                 attachment_paths=attachment_paths)


def generate_rm_spec_record(template_path, spec, parameters,
                            direct_params=None, coa_params=None,
                            completion_fields=None,
                            attachment_paths=None):
    """Generate an RM Specification Test Record with spec info filled in.

    If direct_params/coa_params are provided, they are used for the two
    tables respectively.  Otherwise all parameters go into Table 0.
    """
    return _generate_spec_record(
        template_path, spec, parameters,
        direct_params=direct_params, coa_params=coa_params,
        completion_fields=completion_fields,
        attachment_paths=attachment_paths,
    )


def _generate_spec_record(template_path, spec, parameters,
                          direct_params=None, coa_params=None,
                          completion_fields=None,
                          attachment_paths=None):
    """Shared logic for PK and RM spec record generation.

    Fills header fields (NO, Component No, Rev No, Component Name) and
    vendor line by surgically targeting the correct runs.
    Populates the Characteristic / Specifications table(s).
    Leaves Results / P / F columns blank for hand-filling.
    """
    doc = open_dotx(template_path)
    completion_fields = completion_fields or {}
    lot_number = (completion_fields.get("lot_number") or "").strip()

    # ── Fill header paragraphs ──────────────────────────────────
    for section in doc.sections:
        hdr = section.header
        if not hdr:
            continue
        for p in hdr.paragraphs:
            text = p.text
            p_elem = p._element

            # Check for FILLIN merge fields in this paragraph
            has_fields = bool(p_elem.findall(f".//{{{W_NS}}}fldChar"))

            if "NO.:" in text and "____________" in text and lot_number:
                # "LOT NO.: ____________" — optional fill from print options
                # Leave underscores intact when lot_number is blank.
                for run in p.runs:
                    if "____________" in run.text:
                        _replace_underscores(run, lot_number)

            elif "Component No.:" in text and "Rev. No.:" in text:
                if has_fields:
                    # Strip FILLIN field codes and replace with values
                    _strip_fillin_fields(p, {
                        "Component No.": _val(spec, "material_code"),
                        "Rev. No.": _val(spec, "revision", "00"),
                    })
                else:
                    _fill_after_label(p, "Component No.:", _val(spec, "material_code"))
                    _fill_after_label(p, "Rev. No.:", _val(spec, "revision", "00"))

            elif "Component Name:" in text:
                if has_fields:
                    _strip_fillin_fields(p, {
                        "Component Name": spec["material_name"],
                    })
                else:
                    _fill_after_label(p, "Component Name:", spec["material_name"])

    # ── Fill vendor in body (replace underscores in the correct runs) ──
    supplier_val = (_val(spec, "supplier") or "").strip()
    for p in doc.paragraphs:
        if "Vendor:" not in p.text:
            continue
        for run in p.runs:
            if "Vendor:" in run.text and "___" in run.text and supplier_val:
                _replace_underscores(run, supplier_val)
            # Don't touch "Vendor's Lot No." underscores — filled by hand

    # ── Optional completion fields in body (written/approved/date) ───────
    for p in doc.paragraphs:
        text = p.text
        if "Written By:" in text and "Approved By:" in text:
            _replace_underscores_sequential(p, [
                (completion_fields.get("written_by") or "").strip(),
                (completion_fields.get("written_date") or "").strip(),
                (completion_fields.get("approved_by") or "").strip(),
                (completion_fields.get("approved_date") or "").strip(),
            ])
            _bold_labels_in_paragraph(
                p,
                ["Written By:", "Date:", "Approved By:"]
            )

        elif "Sample Received:" in text and "Logged in By:" in text:
            _bold_labels_in_paragraph(
                p,
                ["Sample Received:", "Time:", "Date:", "Logged in By:"]
            )

        elif "Vendor:" in text and "Lot No.:" in text:
            _bold_labels_in_paragraph(
                p,
                ["Vendor:", "Vendor's Lot No.:", "Vendor’s Lot No.:", "Lot No.:"]
            )

    # ── Fill test parameters table(s) ───────────────────────────
    tables = doc.tables
    if direct_params is not None or coa_params is not None:
        # RM with split parameters
        for table in tables:
            header_text = table.rows[0].cells[0].text.strip()
            if header_text == "Characteristic":
                _fill_spec_table(table, direct_params or [])
            elif "From C of A" in header_text:
                _fill_spec_table(table, coa_params or [])
    else:
        # PK (single table) or RM without split
        for table in tables:
            _fill_spec_table(table, parameters)

    # ── Append 3rd-party attachments as separate pages ────────────
    if attachment_paths:
        _append_attachments(doc, attachment_paths)

    # ── Save ────────────────────────────────────────────────────
    output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(output.name)
    return output.name


def _fill_spec_table(table, parameters):
    """Fill a Characteristic / Specifications table with parameter rows.

    Strategy:
      1. Deep-copy row 1 XML as the formatting template.
      2. Remove all non-header rows from the table.
      3. Append one cloned row per parameter in order.
    This preserves the exact row/cell formatting and keeps correct order.
    """
    if len(table.rows) < 2:
        return

    # Verify this is a spec table
    h0 = table.rows[0].cells[0].text.strip()
    if "Characteristic" not in h0 and "From C of A" not in h0:
        return

    # Save template row XML
    template_tr = deepcopy(table.rows[1]._tr)

    # Remove all non-header rows (iterate in reverse so indices stay valid)
    tbl_elem = table._tbl
    for ri in range(len(table.rows) - 1, 0, -1):
        tbl_elem.remove(table.rows[ri]._tr)

    if not parameters:
        # Keep one empty row so the table isn't just a header
        tbl_elem.append(deepcopy(template_tr))
        return

    # Append one row per parameter
    for param in parameters:
        new_tr = deepcopy(template_tr)
        _set_row_cell_text(new_tr, 0, param["parameter_name"])
        _set_row_cell_text(new_tr, 1, _val(param, "acceptance_criteria"))
        # Cols 2-3 (Results, Reference) left blank for hand-fill
        # Cols 4-5 (P, F) keep their template text for circling on paper
        tbl_elem.append(new_tr)


# ── Attachment appender ─────────────────────────────────────────

def _append_attachments(doc, attachment_paths):
    """Append 3rd-party attachments as separate pages in the document.

    attachment_paths: list of (original_name, file_path) tuples.

    - Images (.png, .jpg, .jpeg) are inserted full-width on a new page.
    - .docx files are merged paragraph-by-paragraph on a new page.
    - .pdf / other: a reference page is added noting the attached file.
    """
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_BREAK

    for original_name, file_path in attachment_paths:
        ext = os.path.splitext(file_path)[1].lower()

        # Add page break before each attachment
        bp = doc.add_paragraph()
        bp.runs[0].add_break(WD_BREAK.PAGE) if bp.runs else bp.add_run().add_break(WD_BREAK.PAGE)

        if ext in (".png", ".jpg", ".jpeg"):
            # Insert image — fit to page width
            heading = doc.add_paragraph()
            run = heading.add_run(f"Attachment: {original_name}")
            run.bold = True
            run.font.size = Pt(11)
            try:
                doc.add_picture(file_path, width=Inches(6.5))
            except Exception:
                doc.add_paragraph(f"[Image could not be embedded: {original_name}]")

        elif ext == ".docx":
            # Merge .docx content paragraph-by-paragraph
            heading = doc.add_paragraph()
            run = heading.add_run(f"Attachment: {original_name}")
            run.bold = True
            run.font.size = Pt(11)
            try:
                att_doc = Document(file_path)
                for p in att_doc.paragraphs:
                    new_p = doc.add_paragraph(p.text, style=p.style.name if p.style else None)
                    # Copy run-level formatting
                    if p.runs and new_p.runs:
                        for src_run, dst_run in zip(p.runs, new_p.runs):
                            if src_run.bold is not None:
                                dst_run.bold = src_run.bold
                            if src_run.italic is not None:
                                dst_run.italic = src_run.italic
                # Copy tables from attached doc
                for table in att_doc.tables:
                    _copy_table(doc, table)
            except Exception:
                doc.add_paragraph(f"[Document could not be embedded: {original_name}]")

        elif ext == ".pdf":
            heading = doc.add_paragraph()
            run = heading.add_run(f"Attachment: {original_name}")
            run.bold = True
            run.font.size = Pt(11)
            doc.add_paragraph(
                "This PDF attachment is included as a separate file. "
                "Please refer to the uploaded PDF document."
            )

        else:
            heading = doc.add_paragraph()
            run = heading.add_run(f"Attachment: {original_name}")
            run.bold = True
            run.font.size = Pt(11)
            doc.add_paragraph(f"[File type {ext} — see uploaded attachment]")


def _copy_table(doc, src_table):
    """Copy a table from one Document into another, preserving structure."""
    rows = len(src_table.rows)
    cols = len(src_table.columns)
    new_table = doc.add_table(rows=rows, cols=cols)
    for ri, row in enumerate(src_table.rows):
        for ci, cell in enumerate(row.cells):
            new_table.rows[ri].cells[ci].text = cell.text


# ══════════════════════════════════════════════════════════════════
# COMPONENT RECEIVING RECORD
# ══════════════════════════════════════════════════════════════════

def generate_receiving_record(template_path, spec, spec_type, po_number=""):
    """Generate a Component Receiving Record from the CLB003 template.

    The template has two identical copies on one page; we keep only one copy.
    Fields filled from the spec:
      - COMPONENT CODE NO.  → spec.material_code
      - COMPONENT NAME      → spec.material_name
      - PO NO.              → po_number (optional)
      - VENDOR              → spec.supplier
    All other fields (LOT NO., PO NO., DATE, etc.) are left blank for hand-fill.

    spec_type: 'pk' or 'rm' (used for filename context only).
    """
    doc = Document(template_path)

    # Keep only one receiving record copy on the sheet
    _remove_second_receiving_copy(doc)

    field_map = {
        "COMPONENT CODE NO.:": _val(spec, "material_code"),
        "COMPONENT NAME:": _val(spec, "material_name"),
        "PO NO.": po_number,
        "VENDOR:": _val(spec, "supplier"),
    }

    for p in doc.paragraphs:
        text = p.text
        for label, value in field_map.items():
            if label in text:
                _fill_underscore_field(p, label, value)

        if "COMPONENT CODE NO.:" in text and "COMPONENT NAME:" in text:
            _bold_labels_in_paragraph(
                p,
                ["COMPONENT CODE NO.:", "COMPONENT NAME:", "LOT NO.:"]
            )
        elif "PO NO.:" in text and "VENDOR:" in text:
            _bold_labels_in_paragraph(
                p,
                ["PO NO.:", "VENDOR:", "MANUFACTURER:", "MFR. LOT NO.:"]
            )

    output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(output.name)
    return output.name


def _remove_second_receiving_copy(doc):
    """Remove duplicate receiving record copy; keep only first form on page."""
    title_indices = [
        i for i, p in enumerate(doc.paragraphs)
        if "COMPONENT RECEIVING RECORD" in p.text.upper()
    ]
    if len(title_indices) < 2:
        return

    second_start = title_indices[1]
    to_remove = list(doc.paragraphs[second_start:])
    for p in to_remove:
        p._element.getparent().remove(p._element)


def _fill_underscore_field(paragraph, label, value):
    """Replace the underscore placeholder after a label within a run's text.

    The template puts multiple fields in a single run, e.g.:
      'COMPONENT CODE NO.: __________  COMPONENT NAME: _____...'
    This function replaces only the underscore block immediately after the
    specified label, leaving other underscore blocks intact.
    """
    if not value:
        return

    for run in paragraph.runs:
        if label in run.text:
            # Find the label position and the underscore block after it
            idx = run.text.index(label) + len(label)
            rest = run.text[idx:]
            m = re.search(r"_{3,}", rest)
            if m:
                before = run.text[:idx + m.start()]
                after = run.text[idx + m.end():]
                run.text = before + str(value) + after
            return


def _bold_labels_in_paragraph(paragraph, labels):
    """Bold specified label tokens and keep non-label text regular."""
    text = paragraph.text
    if not text:
        return

    escaped = [re.escape(lbl) for lbl in labels if lbl]
    if not escaped:
        return

    pattern = re.compile("|".join(sorted(escaped, key=len, reverse=True)))
    matches = list(pattern.finditer(text))
    if not matches:
        return

    p_elem = paragraph._element
    for child in list(p_elem):
        if child.tag in (qn("w:r"), qn("w:hyperlink"), qn("w:smartTag")):
            p_elem.remove(child)

    cursor = 0
    for m in matches:
        if m.start() > cursor:
            run = paragraph.add_run(text[cursor:m.start()])
            run.bold = False
        run = paragraph.add_run(m.group(0))
        run.bold = True
        cursor = m.end()

    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        run.bold = False


# ══════════════════════════════════════════════════════════════════
# SOP
# ══════════════════════════════════════════════════════════════════

def generate_sop(template_path, sop, revision_history=None):
    """Generate an SOP from the SOP Temp.dotx template.

    Fills:
      - Header text-box fields: SOP NO., REV NO.
      - Body section headings with content from the SOP record.
    Leaves signature table and footer untouched.
    """
    doc = open_dotx(template_path)

    # ── Fill header drawing text boxes ──────────────────────────
    for section in doc.sections:
        hdr = section.header
        if not hdr:
            continue
        # The SOP header uses wps:txbx text-box shapes (not tables)
        for txbx in hdr._element.iter(f"{{{WPS_NS}}}txbx"):
            for t_elem in txbx.iter(f"{{{W_NS}}}t"):
                if t_elem.text is None:
                    continue
                txt = t_elem.text.strip()
                if txt.startswith("SOP NO.:"):
                    t_elem.text = f"SOP NO.: {_val(sop, 'sop_number')}  "
                elif txt.startswith("REV NO.:"):
                    t_elem.text = f"REV NO.: {_val(sop, 'revision', '00')}"

    # ── Map section headings → content ──────────────────────────
    section_content = {
        "SUBJECT:": _val(sop, "title"),
        "OBJECTIVE:": _val(sop, "purpose"),
        "RESPONSIBILITIES:": _val(sop, "responsibilities"),
        "FREQUENCY:": _val(sop, "scope"),
        "NECESSARY EQUIPMENT:": _val(sop, "equipment_materials"),
        "PROCEDURE:": _val(sop, "procedure_text"),
        "DOCUMENTATION:": _val(sop, "references_text"),
    }

    # ── Fill sections ───────────────────────────────────────────
    for p in list(doc.paragraphs):
        text_upper = p.text.strip().upper()

        for heading, content in section_content.items():
            heading_upper = heading.upper()
            heading_bare = heading_upper.rstrip(": ")

            if text_upper in (heading_upper, heading_bare):
                if heading == "SUBJECT:":
                    # Append title on the same line, in the existing run
                    if p.runs:
                        # Keep "SUBJECT: " label, append title
                        p.runs[0].text = f"SUBJECT: {content}"
                elif content:
                    _insert_paragraphs_after(p, content)
                break

    # ── Save ────────────────────────────────────────────────────
    output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(output.name)
    return output.name


def _insert_paragraphs_after(ref_para, content):
    """Insert content paragraphs after a heading, using Normal style."""
    lines = content.strip().split("\n")

    for line in reversed(lines):
        new_p = ref_para._element.makeelement(qn("w:p"), {})

        # Set Normal style (not the heading's @Level 1)
        pPr = new_p.makeelement(qn("w:pPr"), {})
        pStyle = pPr.makeelement(qn("w:pStyle"), {})
        pStyle.set(qn("w:val"), "Normal")
        pPr.append(pStyle)
        new_p.append(pPr)

        # Add run with text
        r_elem = new_p.makeelement(qn("w:r"), {})
        t_elem = r_elem.makeelement(qn("w:t"), {})
        t_elem.text = line
        t_elem.set(qn("xml:space"), "preserve")
        r_elem.append(t_elem)
        new_p.append(r_elem)

        ref_para._element.addnext(new_p)


# ══════════════════════════════════════════════════════════════════
# COMPONENT ID TAG (CLB010)
# ══════════════════════════════════════════════════════════════════

# The CLB010 template is a 3×3 table (cols 0 & 2 are tag cells, col 1 is
# a spacer).  That gives 6 tags per page.  Each tag cell contains paragraphs
# with underscore placeholders for:
#   P4: Component# ________ Lot# ____________
#   P6: Component Name:____________________
#   P7:       ____________________________________ (name overflow)
#   P8: Ctn# _____of ______Date: ______By____

TAGS_PER_PAGE = 6
_TAG_CELL_ORDER = [(0, 0), (0, 2), (1, 0), (1, 2), (2, 0), (2, 2)]


def generate_id_tags(template_path, form_data, container_count):
    """Generate Component ID Tags from the CLB010 template.

    form_data dict keys:
        component_code, component_name, lot_number, date, by
    container_count: int — number of tags to produce (each gets Ctn# X of Y).

    Returns path to a temporary .docx file.
    """
    import math

    doc = Document(template_path)
    table = doc.tables[0]

    pages_needed = math.ceil(container_count / TAGS_PER_PAGE)

    # Save a pristine copy of the table XML before we touch anything
    template_tbl_xml = deepcopy(table._tbl)

    # Fill the first page (the existing table)
    _fill_tag_page(table, form_data, start_ctn=1,
                   container_count=container_count)

    # Add additional pages if needed
    for page_idx in range(1, pages_needed):
        # Page break paragraph
        bp = doc.add_paragraph()
        run = bp.add_run()
        from docx.enum.text import WD_BREAK
        run.add_break(WD_BREAK.PAGE)

        # Clone the pristine table
        new_tbl = deepcopy(template_tbl_xml)
        doc.element.body.append(new_tbl)

        # python-docx doesn't automatically track appended XML tables,
        # so we wrap it manually to fill it
        from docx.table import Table
        new_table = Table(new_tbl, doc)

        start_ctn = page_idx * TAGS_PER_PAGE + 1
        _fill_tag_page(new_table, form_data, start_ctn=start_ctn,
                       container_count=container_count)

    output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(output.name)
    return output.name


def _fill_tag_page(table, form_data, start_ctn, container_count):
    """Fill up to 6 tag cells on one page of the table."""
    component_code = form_data.get("component_code", "")
    component_name = form_data.get("component_name", "")
    lot_number = form_data.get("lot_number", "")
    date_val = form_data.get("date", "")
    by_val = form_data.get("by", "")

    for slot_idx, (ri, ci) in enumerate(_TAG_CELL_ORDER):
        ctn_num = start_ctn + slot_idx
        if ctn_num > container_count:
            break  # Leave remaining cells as blank templates

        cell = table.rows[ri].cells[ci]
        _fill_tag_cell(cell, component_code, component_name, lot_number,
                       ctn_num, container_count, date_val, by_val)


def _fill_tag_cell(cell, component_code, component_name, lot_number,
                   ctn_num, ctn_total, date_val, by_val):
    """Fill a single tag cell by replacing underscore placeholders."""
    for p in cell.paragraphs:
        text = p.text

        if "Component#" in text and "Lot#" in text:
            # P4: Component# ________ Lot# ____________
            # Replace the 2 underscore blocks in order: component_code, lot_number
            _replace_underscores_sequential(p, [component_code, lot_number])
            _bold_labels_in_paragraph(p, ["Component#", "Lot#"])

        elif "Component Name:" in text:
            # P6: Component Name:____________________
            if component_name:
                for run in p.runs:
                    if "___" in run.text:
                        run.text = re.sub(r"_{3,}", component_name, run.text, count=1)
                        break
            _bold_labels_in_paragraph(p, ["Component Name:"])

        elif text.strip().startswith("_") and re.search(r"_{10,}", text):
            # P7: continuation underscores — clear only if name was provided
            if component_name:
                for run in p.runs:
                    if "_" in run.text:
                        run.text = re.sub(r"_+", "", run.text)

        elif "Ctn#" in text:
            # P8: Ctn# _____of ______Date: ______By____
            # Replace 4 underscore blocks in order:
            #   1) ctn_num  2) ctn_total  3) date  4) by
            # Pad values with trailing space to preserve readability
            replacements = [
                f"{ctn_num} " if ctn_num else "",
                f"{ctn_total} " if ctn_total else "",
                f"{date_val} " if date_val else "",
                f" {by_val}" if by_val else "",
            ]
            _replace_underscores_sequential(p, replacements)
            _bold_labels_in_paragraph(p, ["Ctn#", "Date:", "By"])


def _replace_underscore_after(paragraph, needle, value, occurrence=1):
    """Replace the Nth underscore block that follows `needle` across runs.

    Because run boundaries vary between cells, we work at the full run
    level — scanning each run for `needle` followed by underscores.
    """
    count = 0
    for run in paragraph.runs:
        # Find all underscore blocks in this run preceded by the needle
        start = 0
        while True:
            idx = run.text.find(needle, start)
            if idx == -1:
                break
            after = idx + len(needle)
            m = re.search(r"_{3,}", run.text[after:])
            if m:
                count += 1
                if count == occurrence:
                    if not value:
                        return  # Leave underscores intact
                    before = run.text[:after + m.start()]
                    rest = run.text[after + m.end():]
                    run.text = before + str(value) + rest
                    return
                start = after + m.end()
            else:
                start = after

    # Fallback: just replace the Nth underscore block overall
    count = 0
    for run in paragraph.runs:
        for m in re.finditer(r"_{3,}", run.text):
            count += 1
            if count == occurrence:
                if not value:
                    return  # Leave underscores intact
                run.text = run.text[:m.start()] + str(value) + run.text[m.end():]
                return


def _replace_underscores_sequential(paragraph, replacements):
    """Replace underscore blocks across all runs in order, one per replacement.

    Empty/blank values leave the underscore block intact so the visual
    underline is preserved for hand-writing.
    """
    rep_idx = 0
    runs = paragraph.runs
    run_idx = 0
    while run_idx < len(runs) and rep_idx < len(replacements):
        run = runs[run_idx]
        offset = 0
        while rep_idx < len(replacements):
            m = re.search(r"_{3,}", run.text[offset:])
            if not m:
                break
            val = replacements[rep_idx]
            rep_idx += 1
            if not val:
                # Empty value — leave underscores intact, skip to next block
                offset = offset + m.end()
                continue
            # Replace this underscore block
            abs_start = offset + m.start()
            abs_end = offset + m.end()
            original_len = len(run.text)
            block_at_end = (abs_end >= original_len)
            run.text = run.text[:abs_start] + str(val) + run.text[abs_end:]
            offset = abs_start + len(str(val))
            # Clear orphan underscores from subsequent runs
            # (handles cases where a single block was split across runs)
            if block_at_end:
                for fi in range(run_idx + 1, len(runs)):
                    orphan_m = re.match(r"^_+", runs[fi].text)
                    if orphan_m:
                        runs[fi].text = runs[fi].text[orphan_m.end():]
                    else:
                        break
        run_idx += 1


# ═══════════════════════════════════════════════════════════════════
# CLB008 – QC Release Stickers (30 per page)
# ═══════════════════════════════════════════════════════════════════

STICKERS_PER_PAGE = 30
_STICKER_CELL_ORDER = [(r, c) for r in range(10) for c in (0, 2, 4)]


def generate_qc_release_stickers(template_path, form_data, sticker_count):
    """Generate CLB008 QC Release Stickers — 30 per page."""
    doc = Document(template_path)
    first_table = doc.tables[0]
    template_tbl_xml = deepcopy(first_table._tbl)

    # Fill the first page (already in the doc)
    _fill_release_page(first_table, form_data, start=1, total=sticker_count)

    # Additional pages as needed
    total_pages = math.ceil(sticker_count / STICKERS_PER_PAGE)
    for page_idx in range(1, total_pages):
        doc.add_page_break()
        new_tbl = deepcopy(template_tbl_xml)
        doc.element.body.append(new_tbl)
        from docx.table import Table
        new_table = Table(new_tbl, doc)
        _fill_release_page(new_table, form_data,
                           start=page_idx * STICKERS_PER_PAGE + 1,
                           total=sticker_count)

    output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(output.name)
    return output.name


def _fill_release_page(table, form_data, start, total):
    """Fill up to 30 QC Release sticker cells on one page."""
    item_val = form_data.get("item_number", "")
    lot_val = form_data.get("lot_number", "")
    date_val = form_data.get("date", "")
    by_val = form_data.get("by", "")

    for slot_idx, (ri, ci) in enumerate(_STICKER_CELL_ORDER):
        sticker_num = start + slot_idx
        if sticker_num > total:
            break
        cell = table.rows[ri].cells[ci]
        _fill_release_cell(cell, item_val, lot_val, date_val, by_val)


def _fill_release_cell(cell, item_val, lot_val, date_val, by_val):
    """Fill a single CLB008 QC Release sticker cell.

    Layout: P0="QC LAB RELEASE", P2="Item#: ___ Lot: ___", P4="Date: ___ By: ___"
    """
    for p in cell.paragraphs:
        text = p.text
        if "Item#" in text and "Lot" in text:
            _replace_underscores_sequential(p, [item_val, lot_val])
            _bold_labels_in_paragraph(p, ["Item#:", "Lot:"])
        elif "Date" in text and "By" in text:
            _replace_underscores_sequential(p, [
                f"{date_val} " if date_val else "",
                f" {by_val}" if by_val else "",
            ])
            _bold_labels_in_paragraph(p, ["Date:", "By:"])


# ═══════════════════════════════════════════════════════════════════
# CLB009 – QC Sampled Stickers (30 per page)
# ═══════════════════════════════════════════════════════════════════

def generate_qc_sampled_stickers(template_path, form_data, sticker_count):
    """Generate CLB009 QC Sampled Stickers — 30 per page."""
    doc = Document(template_path)
    first_table = doc.tables[0]
    template_tbl_xml = deepcopy(first_table._tbl)

    _fill_sampled_page(first_table, form_data, start=1, total=sticker_count)

    total_pages = math.ceil(sticker_count / STICKERS_PER_PAGE)
    for page_idx in range(1, total_pages):
        doc.add_page_break()
        new_tbl = deepcopy(template_tbl_xml)
        doc.element.body.append(new_tbl)
        from docx.table import Table
        new_table = Table(new_tbl, doc)
        _fill_sampled_page(new_table, form_data,
                           start=page_idx * STICKERS_PER_PAGE + 1,
                           total=sticker_count)

    output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(output.name)
    return output.name


def _fill_sampled_page(table, form_data, start, total):
    """Fill up to 30 QC Sampled sticker cells on one page."""
    date_val = form_data.get("date", "")
    by_val = form_data.get("by", "")

    for slot_idx, (ri, ci) in enumerate(_STICKER_CELL_ORDER):
        sticker_num = start + slot_idx
        if sticker_num > total:
            break
        cell = table.rows[ri].cells[ci]
        _fill_sampled_cell(cell, date_val, by_val)


def _fill_sampled_cell(cell, date_val, by_val):
    """Fill a single CLB009 QC Sampled sticker cell.

    Layout: P0="QC LAB", P1="SAMPLED", P3="Date______By: ________"
    """
    for p in cell.paragraphs:
        text = p.text
        if "Date" in text and "By" in text:
            _replace_underscores_sequential(p, [
                f" {date_val} " if date_val else "",
                f" {by_val}" if by_val else "",
            ])
            _bold_labels_in_paragraph(p, ["Date", "By:"])
