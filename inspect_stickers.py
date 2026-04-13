"""Inspect CLB008 and CLB009 sticker template structures."""
from docx import Document
from docx.oxml.ns import qn

for fname in ["CLB008 QC Release Sticker.docx", "CLB009 QC Sampled.docx"]:
    print(f"\n{'='*60}")
    print(f"  {fname}")
    print(f"{'='*60}")

    doc = Document(f"doc_templates/{fname}")

    for si, section in enumerate(doc.sections):
        pw = section.page_width
        ph = section.page_height
        print(f"\nSection {si}: {pw/914400:.2f}in x {ph/914400:.2f}in")
        print(f"  Margins: top={section.top_margin} bot={section.bottom_margin} left={section.left_margin} right={section.right_margin}")

    print(f"\nBody paragraphs: {len(doc.paragraphs)}")
    for pi, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"  P{pi}: text={repr(p.text)}")
            for ri, r in enumerate(p.runs):
                print(f"    R{ri}: text={repr(r.text)} bold={r.bold} sz={r.font.size}")

    print(f"\nTables: {len(doc.tables)}")
    for ti, table in enumerate(doc.tables):
        print(f"\n  Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")

        # Row heights
        for ri, row in enumerate(table.rows):
            tr = row._tr
            trPr = tr.find(qn("w:trPr"))
            h_info = ""
            if trPr is not None:
                trH = trPr.find(qn("w:trHeight"))
                if trH is not None:
                    h_info = f" height={trH.get(qn('w:val'))} rule={trH.get(qn('w:hRule'))}"

            cells_summary = []
            for ci, cell in enumerate(row.cells):
                txt = cell.text.strip().replace("\n", " | ")
                has_pict = bool(cell._tc.findall(".//" + qn("w:pict")))
                has_drawing = bool(cell._tc.findall(".//" + qn("w:drawing")))
                img = " [IMG]" if (has_pict or has_drawing) else ""
                cells_summary.append(f"[{ci}]={repr(txt[:80])}{img}" if txt else f"[{ci}]=<empty>{img}")
            print(f"    Row {ri}{h_info}: {'  '.join(cells_summary)}")

        # Column widths from first row
        print(f"\n    Col widths (row 0):")
        for ci, cell in enumerate(table.rows[0].cells):
            tcPr = cell._tc.find(qn("w:tcPr"))
            if tcPr is not None:
                w = tcPr.find(qn("w:tcW"))
                if w is not None:
                    print(f"      Col {ci}: width={w.get(qn('w:w'))} type={w.get(qn('w:type'))}")

        # Detailed cell paragraphs for first tag cell
        print(f"\n    Detailed cell [0,0] paragraphs:")
        cell = table.rows[0].cells[0]
        for pi, p in enumerate(cell.paragraphs):
            has_img = bool(p._element.findall(".//" + qn("w:pict"))) or bool(p._element.findall(".//" + qn("w:drawing")))
            img_flag = " [IMG]" if has_img else ""
            if p.text.strip() or has_img:
                print(f"      P{pi}: text={repr(p.text)}{img_flag}")
                for rri, r in enumerate(p.runs):
                    print(f"        R{rri}: text={repr(r.text)} bold={r.bold} sz={r.font.size}")
            else:
                print(f"      P{pi}: <empty>")
