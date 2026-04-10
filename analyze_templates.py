#!/usr/bin/env python3
"""
Template Analyzer for Club 13 GMP Toolkit

Analyzes .dotx template files and outputs their structure as JSON.
This helps map where data should be injected by the document generator.

Usage:
    python analyze_templates.py [template_path]
    python analyze_templates.py  # analyzes all templates in doc_templates/
"""
import json
import os
import sys

from docx import Document


def analyze_template(path):
    """Analyze a .dotx template and return its structure."""
    doc = Document(path)
    result = {
        "file": os.path.basename(path),
        "file_size": os.path.getsize(path),
        "sections": [],
        "paragraphs": [],
        "tables": [],
        "headers": [],
        "footers": [],
        "styles_used": set(),
        "potential_fields": [],
    }

    # Analyze paragraphs
    for i, p in enumerate(doc.paragraphs):
        style_name = p.style.name if p.style else "None"
        result["styles_used"].add(style_name)
        para_info = {
            "index": i,
            "text": p.text,
            "style": style_name,
            "alignment": str(p.alignment) if p.alignment else None,
            "runs": [],
        }
        for run in p.runs:
            para_info["runs"].append({
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "font_name": run.font.name,
                "font_size": str(run.font.size) if run.font.size else None,
            })

        # Check for potential fillable fields
        text = p.text
        if any(marker in text for marker in ["{{", "}}", "<<", ">>", "«", "»", "[", "]", "____", "___"]):
            result["potential_fields"].append({"location": f"paragraph_{i}", "text": text})
        if ":" in text and len(text.split(":")) == 2:
            label = text.split(":")[0].strip()
            value = text.split(":")[1].strip()
            if len(label) < 40 and (not value or value.startswith("_") or value.startswith(" ")):
                result["potential_fields"].append({"location": f"paragraph_{i}", "label": label, "text": text})

        result["paragraphs"].append(para_info)

    # Analyze tables
    for i, table in enumerate(doc.tables):
        table_info = {
            "index": i,
            "num_rows": len(table.rows),
            "num_cols": len(table.columns) if table.rows else 0,
            "rows": [],
        }
        for j, row in enumerate(table.rows):
            row_cells = []
            for cell in row.cells:
                cell_info = {
                    "text": cell.text,
                    "paragraphs": [p.text for p in cell.paragraphs],
                }
                row_cells.append(cell_info)
                # Check for potential fields in table cells
                if any(marker in cell.text for marker in ["{{", "}}", "<<", ">>", "«", "»", "____", "___"]):
                    result["potential_fields"].append({"location": f"table_{i}_row_{j}", "text": cell.text})
                if not cell.text.strip() and j > 0:
                    # Empty cell that might be a fill point
                    pass
            table_info["rows"].append(row_cells)
        result["tables"].append(table_info)

    # Analyze headers and footers
    for si, section in enumerate(doc.sections):
        section_info = {
            "index": si,
            "page_width": str(section.page_width),
            "page_height": str(section.page_height),
            "left_margin": str(section.left_margin),
            "right_margin": str(section.right_margin),
            "top_margin": str(section.top_margin),
            "bottom_margin": str(section.bottom_margin),
        }
        result["sections"].append(section_info)

        # Header
        try:
            if section.header and not section.header.is_linked_to_previous:
                header_info = {"section": si, "paragraphs": [], "tables": []}
                for p in section.header.paragraphs:
                    header_info["paragraphs"].append({"text": p.text, "style": p.style.name if p.style else None})
                for table in section.header.tables:
                    rows = []
                    for row in table.rows:
                        rows.append([cell.text for cell in row.cells])
                    header_info["tables"].append(rows)
                result["headers"].append(header_info)
        except Exception as e:
            result["headers"].append({"section": si, "error": str(e)})

        # Footer
        try:
            if section.footer and not section.footer.is_linked_to_previous:
                footer_info = {"section": si, "paragraphs": [], "tables": []}
                for p in section.footer.paragraphs:
                    footer_info["paragraphs"].append({"text": p.text, "style": p.style.name if p.style else None})
                for table in section.footer.tables:
                    rows = []
                    for row in table.rows:
                        rows.append([cell.text for cell in row.cells])
                    footer_info["tables"].append(rows)
                result["footers"].append(footer_info)
        except Exception as e:
            result["footers"].append({"section": si, "error": str(e)})

    result["styles_used"] = sorted(list(result["styles_used"]))
    return result


def main():
    if len(sys.argv) > 1:
        paths = sys.argv[1:]
    else:
        template_dir = os.path.join(os.path.dirname(__file__), "doc_templates")
        if not os.path.exists(template_dir):
            # Check production location
            if os.path.exists("/data/doc_templates"):
                template_dir = "/data/doc_templates"
            else:
                print(f"Template directory not found: {template_dir}")
                print("Usage: python analyze_templates.py [template_path]")
                sys.exit(1)
        paths = [os.path.join(template_dir, f) for f in os.listdir(template_dir) if f.endswith((".dotx", ".docx"))]

    if not paths:
        print("No template files found.")
        sys.exit(1)

    for path in paths:
        if not os.path.exists(path):
            print(f"File not found: {path}")
            continue
        print(f"\n{'='*60}")
        print(f"Analyzing: {os.path.basename(path)}")
        print(f"{'='*60}")

        result = analyze_template(path)

        print(f"\nFile size: {result['file_size']:,} bytes")
        print(f"Sections: {len(result['sections'])}")
        print(f"Paragraphs: {len(result['paragraphs'])}")
        print(f"Tables: {len(result['tables'])}")
        print(f"Headers: {len(result['headers'])}")
        print(f"Footers: {len(result['footers'])}")
        print(f"Styles used: {', '.join(result['styles_used'])}")

        if result["sections"]:
            s = result["sections"][0]
            print(f"\nPage dimensions:")
            print(f"  Width: {s['page_width']}, Height: {s['page_height']}")
            print(f"  Margins: L={s['left_margin']} R={s['right_margin']} T={s['top_margin']} B={s['bottom_margin']}")

        if result["potential_fields"]:
            print(f"\nPotential fillable fields ({len(result['potential_fields'])}):")
            for field in result["potential_fields"]:
                print(f"  [{field['location']}] {field.get('label', field.get('text', ''))[:80]}")

        if result["tables"]:
            print(f"\nTables:")
            for table in result["tables"]:
                print(f"  Table {table['index']}: {table['num_rows']} rows x {table['num_cols']} cols")
                if table["rows"]:
                    # Show first row (header)
                    header = [cell["text"][:30] for cell in table["rows"][0]]
                    print(f"    Header: {' | '.join(header)}")

        if result["headers"]:
            print(f"\nHeaders:")
            for h in result["headers"]:
                if "error" in h:
                    print(f"  Section {h['section']}: Error - {h['error']}")
                else:
                    for p in h.get("paragraphs", []):
                        if p["text"].strip():
                            print(f"  [{p['style']}] {p['text'][:80]}")

        if result["footers"]:
            print(f"\nFooters:")
            for f in result["footers"]:
                if "error" in f:
                    print(f"  Section {f['section']}: Error - {f['error']}")
                else:
                    for p in f.get("paragraphs", []):
                        if p["text"].strip():
                            print(f"  [{p['style']}] {p['text'][:80]}")

        # Save full JSON
        json_path = path.rsplit(".", 1)[0] + "_analysis.json"
        with open(json_path, "w") as f:
            json.dump(result, f, indent=2, default=str)
        print(f"\nFull analysis saved to: {json_path}")


if __name__ == "__main__":
    main()
