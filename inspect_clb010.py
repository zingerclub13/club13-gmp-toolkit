"""Inspect CLB010 Component ID Tag template structure — Phase 2."""
from docx import Document
from docx.oxml.ns import qn

doc = Document("doc_templates/CLB010 Component ID Tag.docx")
t = doc.tables[0]

# Column widths
print("=== COLUMN WIDTHS (Row 0) ===")
for ci, cell in enumerate(t.rows[0].cells):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is not None:
        w = tcPr.find(qn("w:tcW"))
        if w is not None:
            wval = w.get(qn("w:w"))
            wtype = w.get(qn("w:type"))
            print(f"  Col {ci}: width={wval} type={wtype}")

# Table borders
print("\n=== TABLE BORDERS ===")
tblPr = t._tbl.find(qn("w:tblPr"))
if tblPr is not None:
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is not None:
        for b in borders:
            tag = b.tag.split("}")[-1] if "}" in b.tag else b.tag
            val = b.get(qn("w:val"))
            sz = b.get(qn("w:sz"))
            print(f"  {tag}: val={val} sz={sz}")

# Cell borders
print("\n=== CELL BORDERS (Row 0) ===")
for ci, cell in enumerate(t.rows[0].cells):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is not None:
        borders = tcPr.find(qn("w:tcBorders"))
        if borders is not None:
            parts = []
            for b in borders:
                tag = b.tag.split("}")[-1] if "}" in b.tag else b.tag
                val = b.get(qn("w:val"))
                parts.append(f"{tag}={val}")
            print(f"  Col {ci}: {', '.join(parts)}")
        else:
            print(f"  Col {ci}: no cell borders")

# Check for image/drawing in each cell
print("\n=== IMAGES/DRAWINGS PER CELL ===")
for ri, row in enumerate(t.rows):
    for ci, cell in enumerate(row.cells):
        for pi, p in enumerate(cell.paragraphs):
            drawings = p._element.findall(".//" + qn("w:drawing"))
            picts = p._element.findall(".//" + qn("w:pict"))
            if drawings:
                print(f"  [{ri},{ci}] P{pi}: {len(drawings)} drawing(s)")
            if picts:
                print(f"  [{ri},{ci}] P{pi}: {len(picts)} pict(s)")

# Full paragraph listing for cell [0,0]
print("\n=== ALL PARAGRAPHS IN CELL [0,0] (incl empty) ===")
cell = t.rows[0].cells[0]
for pi, p in enumerate(cell.paragraphs):
    has_drawing = bool(p._element.findall(".//" + qn("w:drawing")))
    print(f"  P{pi}: text={p.text!r} drawing={has_drawing} runs={len(p.runs)}")
