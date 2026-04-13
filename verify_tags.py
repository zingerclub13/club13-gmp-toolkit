"""Verify tag generation output."""
from docx import Document

doc = Document("/tmp/test_tags_multi.docx")
print(f"Tables: {len(doc.tables)}")
order = [(0, 0), (0, 2), (1, 0), (1, 2), (2, 0), (2, 2)]
for ti, t in enumerate(doc.tables):
    print(f"\n=== Page {ti+1} ===")
    for slot, (ri, ci) in enumerate(order):
        cell = t.rows[ri].cells[ci]
        ctn_line = ""
        for p in cell.paragraphs:
            if "Ctn#" in p.text or "Component#" in p.text:
                ctn_line += p.text.strip() + "  "
        tag_num = ti * 6 + slot + 1
        print(f"  Tag {tag_num}: {ctn_line.strip()}")
