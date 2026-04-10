"""Check SOP header structure."""
import zipfile
import io
from docx import Document


def open_dotx(path):
    with open(path, "rb") as f:
        data = io.BytesIO(f.read())
    patched = io.BytesIO()
    with zipfile.ZipFile(data, "r") as zin, zipfile.ZipFile(patched, "w") as zout:
        for item in zin.infolist():
            content = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                content = content.replace(
                    b"wordprocessingml.template.main+xml",
                    b"wordprocessingml.document.main+xml",
                )
            zout.writestr(item, content)
    patched.seek(0)
    return Document(patched)


doc = open_dotx("doc_templates/SOP Temp.dotx")
for si, section in enumerate(doc.sections):
    hdr = section.header
    if hdr:
        for p in hdr.paragraphs:
            if p.text.strip():
                print(f"Header P: {repr(p.text)}")
        for ht, htable in enumerate(hdr.tables):
            print(f"Header Table {ht}: {len(htable.rows)}x{len(htable.columns)}")
            for ri, row in enumerate(htable.rows):
                cells = [c.text.strip()[:60] for c in row.cells]
                unique = []
                prev = None
                for c in cells:
                    if c != prev:
                        unique.append(c)
                    prev = c
                print(f"  Row {ri}: {unique}")

# Also check PK header tables
print("\n--- PK Header ---")
doc2 = open_dotx("doc_templates/PK Specification Test Record Template.dotx")
for si, section in enumerate(doc2.sections):
    hdr = section.header
    if hdr:
        for ht, htable in enumerate(hdr.tables):
            print(f"Header Table {ht}: {len(htable.rows)}x{len(htable.columns)}")
            for ri, row in enumerate(htable.rows):
                cells = [c.text.strip()[:60] for c in row.cells]
                unique = []
                prev = None
                for c in cells:
                    if c != prev:
                        unique.append(c)
                    prev = c
                print(f"  Row {ri}: {unique}")
        for p in hdr.paragraphs:
            # show run-level detail
            runs = [(r.text, r.bold, r.font.size) for r in p.runs if r.text.strip()]
            if runs:
                print(f"  P: {runs}")
