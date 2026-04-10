"""Analyze the .dotx templates to understand their structure."""
import shutil
import tempfile
from docx import Document

templates = [
    "PK Specification Test Record Template.dotx",
    "RM Specification Test Record Template.dotx",
    "SOP Temp.dotx",
]


def open_dotx(path):
    """Open a .dotx file by patching its content type to .docx."""
    import zipfile
    import io
    import re

    with open(path, "rb") as f:
        data = io.BytesIO(f.read())

    # Read zip, patch [Content_Types].xml
    patched = io.BytesIO()
    with zipfile.ZipFile(data, "r") as zin, zipfile.ZipFile(patched, "w") as zout:
        for item in zin.infolist():
            content = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                content = content.replace(
                    b"wordprocessingml.template.main+xml",
                    b"wordprocessingml.document.main+xml",
                )
            zout.writestr(item, content)

    patched.seek(0)
    return Document(patched)


for name in templates:
    print(f"\n{'='*60}")
    print(f"  {name}")
    print(f"{'='*60}")
    doc = open_dotx(f"doc_templates/{name}")

    print(f"\nParagraphs: {len(doc.paragraphs)}")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"  P{i}: [{p.style.name}] {repr(p.text[:120])}")

    print(f"\nTables: {len(doc.tables)}")
    for ti, t in enumerate(doc.tables):
        print(f"\n  Table {ti}: {len(t.rows)} rows x {len(t.columns)} cols")
        for ri, row in enumerate(t.rows):
            cells = [c.text.strip()[:50] for c in row.cells]
            unique = []
            prev = None
            for c in cells:
                if c != prev:
                    unique.append(c)
                prev = c
            print(f"    Row {ri}: {unique}")
            if ri > 25:
                print(f"    ... ({len(t.rows) - ri} more rows)")
                break

    # Check headers/footers
    for si, section in enumerate(doc.sections):
        hdr = section.header
        if hdr and hdr.paragraphs:
            htexts = [p.text for p in hdr.paragraphs if p.text.strip()]
            if htexts:
                print(f"\n  Header (section {si}): {htexts}")
            for ht, htable in enumerate(hdr.tables):
                print(f"  Header Table {ht}: {len(htable.rows)}x{len(htable.columns)}")
                for ri, row in enumerate(htable.rows):
                    cells = [c.text.strip()[:50] for c in row.cells]
                    unique = []
                    prev = None
                    for c in cells:
                        if c != prev:
                            unique.append(c)
                        prev = c
                    print(f"    Row {ri}: {unique}")
        ftr = section.footer
        if ftr and ftr.paragraphs:
            ftexts = [p.text for p in ftr.paragraphs if p.text.strip()]
            if ftexts:
                print(f"\n  Footer (section {si}): {ftexts}")
